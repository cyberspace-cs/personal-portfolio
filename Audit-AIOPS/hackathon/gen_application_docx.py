# -*- coding: utf-8 -*-
"""生成 AdventureX 黑客松「申请更新材料」DOCX 附件（v2：更具体、去 AI 腔）。

输出：hackathon/AdventureX-申请更新材料.docx
【】内为占位/举例，请用 Word 打开替换真实信息后再上传。
"""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
INK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)
ACCENT = RGBColor(0x2E, 0x6B, 0xA8)
CN_FONT = "微软雅黑"
OUT = os.path.join(os.path.dirname(__file__), "AdventureX-申请更新材料.docx")


def set_cn_font(run, font=CN_FONT):
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_para(doc, text="", *, size=10.5, bold=False, color=INK,
             align=None, before=6, after=6, line=1.5, font=CN_FONT, italic=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        set_cn_font(run, font)
    return p


def add_title(doc, text):
    return add_para(doc, text, size=20, bold=True, color=NAVY,
                    align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2, line=1.2)


def add_subtitle(doc, text):
    return add_para(doc, text, size=11, color=GRAY,
                    align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=14, line=1.2)


def add_h2(doc, text):
    return add_para(doc, text, size=13.5, bold=True, color=NAVY,
                    before=14, after=6, line=1.3)


def add_body(doc, text):
    return add_para(doc, text, size=10.5, color=INK,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=6, line=1.6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK
    set_cn_font(run)
    return p


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.4); s.right_margin = Cm(2.4)
    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)

    add_title(doc, "AdventureX 黑客松 · 申请更新材料")
    add_subtitle(doc, "【姓名】 ｜ 【电话】 ｜ 【邮箱】 ｜ 2026 年 X 月")

    add_h2(doc, "一、这次更新想说什么")
    add_body(doc,
        "最近我做完了一个自己比较满意的项目——Audit-AIOPS，一个给审计运维场景用的 AI Agent 助手。"
        "它是我下班后一点点攒出来的，不是 demo，是能跑完整条链路的东西：一句话诉求进去，工单拆好、路由好、"
        "还能回答带出处的问题。我想带着它来 AdventureX，看能不能在 48 小时里做出点真能让人点开就用的东西。")

    add_h2(doc, "二、关于我")
    add_body(doc,
        "我小时候不算多顺。高考没考好，去了山东大学自动化专业。那阵子挺丧的，但我没躺平——"
        "大一基本住在自习室，成绩冲到年级前面，成功转到了网络空间安全专业。说实话，转过去之后才觉得"
        "「找对地方了」。大三那年我和队友做了一个信息安全类的作品，拿了国家级一等奖；具体做的是"
        "【例如：针对某类 Web 攻击的检测/溯源系统，或一套内网资产测绘工具】，从需求调研到代码实现都是我们自己啃下来的。")
    add_body(doc,
        "大四我本来能保研浙大网安，但我最后选了考研北大软工。周围人都觉得我亏了——保研多稳啊。"
        "但我想去更接近「把东西做出来、给人用」的环境，不愿意待在舒适区。后来考上了。")
    add_body(doc,
        "现在我在部委计算中心做【例如：某类业务系统的研发 / 安全运维】。工作稳定、体面，但我干了一年多，"
        "越来越觉得缺了点什么——直到去年参加了一场黑客松。48 小时里，从一拍脑袋的想法到一个能点的原型，"
        "那种节奏让我一下子想起大三通宵做作品的感觉。我意识到，这可能才是我一直想要的状态：用一个周末，"
        "逼自己把脑子里的东西真正做出来。")

    add_h2(doc, "三、我做的项目：Audit-AIOPS（为什么我觉得它能证明点什么）")
    add_body(doc,
        "说它「能跑」是有证据的，不是 PPT：")
    add_bullet(doc,
        "用户一句话（比如「我要借台终端还要开个视频会」），系统自动拆成多个工单，分别派给对应责任人、"
        "设好截止时间，不用人工二次分派——解决的是「一句话里多件事、人肉拆容易漏」的真实麻烦。")
    add_bullet(doc,
        "知识问答不瞎编：检索走四路融合——关键词、向量、图 RAG（把审计里「审批/资产/权限」这些实体和它们"
        "的关系建成一张图，能召回「关系相关但字面没出现」的文档）、多模态 RAG（连表格和截图都能被搜到）。"
        "每条回答都带出处，满足审计合规。")
    add_bullet(doc,
        "推理成本是真的压下来了：蒸馏、INT8 量化、剪枝、投机解码、Prompt 缓存，我全在 CPU 上跑通并测了数——"
        "单轮 token 成本从 ¥0.042 降到 ¥4.9e-5，月度能省约 99.9%；前缀缓存命中率 99.75%、首字延迟降 87.7%。"
        "这些都可以 `python sft/xxx.py` 当场复现。")
    add_bullet(doc,
        "涉密内网能部署：模型压缩后在内网 CPU 上就能跑，不依赖外网 API——这是 ToB/政务场景最在意的一点。")
    add_body(doc,
        "我把它做出来，就是为了证明一件事：我还能独立把一套复杂系统从零拉起来，而且拉起来的东西是能用的，"
        "不是花架子。这大概也是我想带进 AdventureX 的最实在的底气。")

    add_h2(doc, "四、来 AdventureX 我想做什么")
    add_body(doc,
        "比起在比赛里再做一个「聊天机器人」，我更想做一件具体的小事。我目前的想法是：把 Audit-AIOPS 里"
        "已经跑通的能力（Agent 自动拆单 + 四路检索 + 低成本推理）搬到一个更日常的场景，做一个面向科研党/"
        "学生的小工具——比如帮人管文献、记实验、自动整理成可检索的知识库，48 小时后别人能真的点开、"
        "真的往里塞东西、真的查到。当然，具体方向也可以根据现场主题和队友再定，我的核心标准只有一个："
        "做出来得是「能用的」，不是「能讲的」。")

    add_h2(doc, "五、我能带来的")
    add_bullet(doc,
        "后端（Python/FastAPI）、前端原型、Agent 编排、RAG、推理优化都能自己搞；一个想法到我手里，"
        "能变成可演示的东西，不依赖别人搭脚手架。")
    add_bullet(doc,
        "一套已经验证可用的技术底座：Agent 编排 + 四路检索 + 推理优化全跑通，黑客松期间我能把时间花在"
        "产品打磨上，而不是从零搭环境。")
    add_bullet(doc,
        "一个被创造欲驱动、也扛得住通宵的人；以及——我真在部委里干过系统研发，知道「能上线」和"
        "「能 demo」之间的距离。")
    add_body(doc, "如果在读这段文字的你也在找队友，欢迎找我。")

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()

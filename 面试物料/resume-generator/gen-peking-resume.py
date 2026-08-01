# -*- coding: utf-8 -*-
"""生成「北大风格」单页式简历（【姓名】），重点突出 coach + AIOPS 双旗舰 AI 项目。

版式语言参考北京大学专属简历模板：姓名大标题 + 求职意向 + 联系方式，
分节（教育/实习/项目/技能/荣誉）用粗体小标题 + 细分割线 + 「·」要点。
页脚保留投递提示（与分厂适配包一致）。

输出：output/【姓名】-简历-北大风格.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "宋体"
FONT_H = "微软雅黑"

doc = Document()

# 页边距
for s in doc.sections:
    s.top_margin = Cm(1.6)
    s.bottom_margin = Cm(1.6)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)
    # 页脚
    footer = s.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("投递前请使用对应分厂适配简历并以官网最新 JD 复核 · 招聘入口见各分厂适配包")
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)

# 默认字体
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_font(run, size=10.5, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font)


def heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    r = p.add_run(text)
    set_font(r, size=12, bold=True, font=FONT_H)
    r.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    # 底部细线
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1e3a8a")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def entry(title, sub, lines):
    p = doc.add_paragraph()
    p.space_before = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True)
    if sub:
        r2 = p.add_run("   " + sub)
        set_font(r2, size=9.5, color=RGBColor(0x55, 0x55, 0x55))
    for ln in lines:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Cm(0.4)
        bp.paragraph_format.space_after = Pt(1)
        br = bp.add_run("· " + ln)
        set_font(br, size=10)


# ===== 头部 =====
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
nr = name_p.add_run("【姓名】")
set_font(nr, size=22, bold=True, font=FONT_H)

obj_p = doc.add_paragraph()
orun = obj_p.add_run("求职意向：大模型算法工程师 / AI Agent 应用开发")
set_font(orun, size=11, bold=True, color=RGBColor(0x25, 0x25, 0x25))

contact = doc.add_paragraph()
crun = contact.add_run("23岁 · 男 · 【电话】 · 【邮箱】 · 中国·北京")
set_font(crun, size=9.5, color=RGBColor(0x44, 0x44, 0x44))

# ===== 教育背景 =====
heading("教育背景")
entry("北京大学  软件工程（硕士）", "2022.09 – 2025.07", [
    "GPA 3.56/4（专业排名前 10%），获北京大学三好学生、九坤奖学金、社会工作奖",
    "主修：算法设计与分析、机器学习、面向对象技术、软件体系结构与设计、Python 语言处理",
])
entry("山东大学  网络空间安全（学士）", "2018.09 – 2022.06", [
    "GPA 3.61/5，获省级大学生信息安全竞赛一等奖、校级优秀毕业生",
    "主修：软件安全、机器学习、数据结构与算法、计算机组成与设计、计算机网络",
])

# ===== 实习经验 =====
heading("实习经验")
entry("昆仑万维  AI 游戏算法实习生", "2023.09 – 2024.02", [
    "提优大模型决策过程：借鉴 LLM+P 思路基于 Python 完成游戏大模型决策规划 coding，token 量降至 3/4、速度提升 2 倍",
    "Prompt Engineering：模板化生成 3D 物体（Animals/Furnitures/Architectures）训练数据",
    "系统化数据生成与分析：用爬虫 + 类型关键词索引构建剧情大模型数据体系，按剧情占比/类型/畅销度分层",
])

# ===== 项目经验（重点：coach + AIOPS 双旗舰）=====
heading("项目经验")
entry("专属刷题教练 / 定制化备考 Agent（独立作者）", "大模型 Agent · 全栈", [
    "多源题库（考研/考公/大厂面试）聚合平台，升级为 LLM 备考 Agent：LangGraph 编排 + 混元/通义双基座",
    "工程能力：MCP 工具调用、分层记忆（SQLite 持久化 + 五段式上下文预算）、RAG 引用溯源、反思 Agent、学习异常检测（AIOps 迁移）",
    "效果：个性化推荐 + 错题复盘 + 进度可视化；无 API Key 可规则降级运行，已在个人网站上线可交互 Demo",
])
entry("Audit-AIOPS：基于大模型的企业审计日志异常检测系统", "大模型 · AIOps", [
    "负责日志模板解析与结构化、基于 LLM 的异常语义检测与根因分析、告警收敛与可解释报告生成",
    "相较规则基线提升异常召回并降低误报，落地内部运维场景，沉淀为「学习异常检测」能力反哺备考 Agent",
])
entry("基于函数加密的隐私保护神经网络系统（队长）", "隐私计算 · 国家级奖项", [
    "提出 FS-CryptoNN 框架，实现训练/预测双阶段完全安全保护；较开源方案节省 33% 内存、提升 5 倍性能",
    "发现并修复两处严重安全漏洞；获 2021 全国大学生信息安全竞赛作品赛决赛一等奖（国家级）",
])

# ===== 技能 =====
heading("专业技能")
sk = doc.add_paragraph()
skr = sk.add_run("编程语言：Python（精通）、Java、C++　|　大模型：PyTorch、LoRA/QLoRA 微调、RAG、Prompt Engineering、LangChain/LangGraph\n"
                 "工程：FastAPI、Docker、vLLM、SQLite　|　Agent：MCP、分层记忆、Context Harness Loop、反思 Agent\n"
                 "其他：隐私计算/函数加密、数据结构与算法、机器学习、审计日志异常检测")
set_font(skr, size=10)

# ===== 荣誉 =====
heading("荣誉证书")
hon = doc.add_paragraph()
hr = hon.add_run("· 2021 全国大学生信息安全竞赛作品赛一等奖（国家级）　· 美国大学生数学建模大赛三等奖\n"
                 "· 北京大学三好学生 / 九坤奖学金 / 社会工作奖　· 山东大学优秀毕业生　· CET-6")
set_font(hr, size=10)

doc.save("output/【姓名】-简历-北大风格.docx")
print("OK -> output/【姓名】-简历-北大风格.docx")
